from time import timezone
import urllib
from urllib.request import Request, urlopen
import requests
from bs4 import BeautifulSoup
import re
from django.utils import timezone
from bs4.element import Comment



import os
from io import BytesIO
import PyPDF2
import docx


import textract

import tempfile
import pythoncom


from .models import URL, Data

import win32com.client


#Removes all occurence of val(a particular url) from userlist(a list of urls)
def remove(val,userlist):
    try:
        while True:
            userlist.remove(val)
    except ValueError:
        pass
    
    return userlist

#Removes all duplicates from a list
def deleteduplicates(userlist):
    return list(dict.fromkeys(userlist))
    




#Takes a url and returns list of all links in that url
def geturls(url):
    try:
        existing_url = URL.objects.get(name=url).neighbours.all()
    except URL.DoesNotExist:
        existing_url = None
    if(existing_url):
        existing_url = URL.objects.get(name=url).neighbours.all()
        existing_url = list(existing_url)
        existing_url.append(URL.objects.get(name=url))
        li=[]
        for i in existing_url:
            li.append(i.name)
        print(li)
        
        return li


    try:
        resp = urllib.request.urlopen(url)
    except:
        return []

    soup = BeautifulSoup(resp, from_encoding=resp.info().get_param('charset'), features="html.parser")  
    external_links = set()
    internal_links = set()
    for line in soup.find_all('a'):
        link = line.get('href')
        if not link:
            continue
        if link.startswith('http'):
            external_links.add(link)
        else:
            internal_links.add(link)

    # Depending on usage, full internal links may be preferred. 
    full_internal_links = {
        urllib.parse.urljoin(url, internal_link) 
        for internal_link in internal_links
    }

    data = []

    # all unique external and full internal links.
    for link in external_links.union(full_internal_links):
        data.append(link)


    data = deleteduplicates(data)

    data = [x for x in data if not x.endswith("pdf")]
    data = [x for x in data if not x.endswith("doc")]


    data = [x for x in data if x.startswith("http")]

    data = remove(url,data)


    #Saving the scraped URL in URL Model
    try:
        url_model = URL.objects.get(name=url)
    except URL.DoesNotExist:
        url_model = None

    if(url_model):
        print('URL Model already exists for: '+url)
    else:    
        url_model = URL()
        url_model.name = url
        url_model.scraped_date = timezone.datetime.now()
        url_model.save()
        print('URL Model created for: '+url)

    #Saving Neighbours of URL
    for i in data:

        try:
            urlneighbour_model = URL.objects.get(name=i)
        except URL.DoesNotExist:
            urlneighbour_model = None

        if(urlneighbour_model):
            url_model.neighbours.add(urlneighbour_model)
        else:
            urlneighbour_model = URL()
            urlneighbour_model.name = i
            urlneighbour_model.scraped_date = timezone.datetime.now()
            urlneighbour_model.save()
            
            urlneighbour_model = URL.objects.get(name=i)
            url_model.neighbours.add(urlneighbour_model)
            
            print(url+' :URL has been added for: '+ i)


    #Getting the 
    obj = URL.objects.get(name=url)
    nbhd = list(obj.neighbours.all())
    nbhd.append(obj)
    
    li=[]
    for i in nbhd:
        if(i is None):
            continue
        li.append(i.name)


    print(li)

    return li




#Takes url and depth, and returns list of all urls
def geturlswdepth(url, depth):
    depth = depth - 1
    finallist = []
    
    if depth==0:
        try:
            url_model = URL.objects.get(name=url)
        except URL.DoesNotExist:
            url_model = None

        if(url_model):
            print('URL Model already exists for: '+url)
        else:    
            url_model = URL()
            url_model.name = url
            url_model.scraped_date = timezone.datetime.now()
            url_model.save()
            print('URL Model created for: '+url)

        finallist.append(url)
        return finallist
    
    if depth ==1:
        finallist = geturls(url)
        return finallist
    if depth>1:
        finallist = geturls(url)
        newlist = finallist

        newlist = finallist[:-1]

        for i in range(depth-1):
            currentlist = []
            for link in newlist:
                templist = geturls(link)
                currentlist.append(templist[:-1]) 
            
            finallist = finallist + currentlist
            newlist = currentlist
        
        return finallist









#Filter for gettext function
def tag_visible(element):
    if element.parent.name in ['style', 'script', 'head', 'title', 'meta', '[document]']:
        return False
    if isinstance(element, Comment):
        return False
    return True

#Takes url then makes model for that url and returns 1 if successful or None
def gettext(url):
    try:
        existing_data = Data.objects.get(data_url=url)
    except Data.DoesNotExist:
        existing_data = None
    if(existing_data):
        return 1

    try:
        r = requests.get(url)
    except:
        return None
    soup = BeautifulSoup(r.content, 'html.parser')
    texts = soup.findAll(text=True)
    visible_texts = filter(tag_visible, texts)  
    
    data = u" ".join(t.strip() for t in visible_texts)

    url_model = URL.objects.get(name=url)

    data_model = Data()
    data_model.data_url = url
    data_model.type = 'text'
    data_model.content = data
    data_model.scraped_date = timezone.datetime.now()
    data_model.save()
    data_model.url.add(url_model)
    data_model.save()



    return 1







#Takes pdf url and returns text of that pdf
def getpdf(url):
    
    try:
        existing_url = URL.objects.get(name=url)
        existing_data = Data.objects.filter(url=existing_url, type='pdf')
    except:
        existing_data = None
    if(existing_data):
        existing_data = Data.objects.filter(url=existing_url, type='pdf').only('data_url')
        existing_data = list(existing_data)
        existing_data.append(URL.objects.get(name=url).name)
        li=[]
        for i in existing_data:
            li.append(i)
        
        return li
    
    allpdf = getlinksoffile(url,'pdf')
    scraped_pdf = []
    for i in allpdf:
        
        data = givepdf(i)

        url_model = URL.objects.get(name=url)

        data_model = Data()
        print("PDF is "+ i)
        data_model.data_url = i
        data_model.type = 'pdf'
        data_model.content = data
        data_model.scraped_date = timezone.datetime.now()
        data_model.save()
        data_model.url.add(url_model)
        data_model.save()


        scraped_pdf.append(i)

    
    return scraped_pdf






#Takes docx url and reuturns text of that docx
def getword(url):

    try:
        existing_url = URL.objects.get(name=url)
        existing_data = Data.objects.filter(url=existing_url, type='doc')
    except:
        existing_data = None
    if(existing_data):
        existing_data = Data.objects.filter(url=existing_url, type='doc').only('data_url')
        existing_data = list(existing_data)
        existing_data.append(URL.objects.get(name=url).name)
        li=[]
        for i in existing_data:
            li.append(i)
        
        return li

    alldoc = getlinksoffile(url,'doc')
    print("Links gotten from getlinksfile: ")
    print(alldoc)
    scraped_doc = []
    for i in alldoc:

        data = givedocx(i)

        url_model = URL.objects.get(name=url)

        data_model = Data()
        data_model.data_url = i
        data_model.type = 'doc'
        data_model.content = data
        data_model.scraped_date = timezone.datetime.now()
        data_model.save()
        data_model.url.add(url_model)
        data_model.save()


        scraped_doc.append(i)

    return scraped_doc

    






#Takes a url and type, and returns a list of all url links or that datatype
def getlinksoffile(url,type):
    try:
        resp = urllib.request.urlopen(url)
    except:
        return []
    # Get server encoding per recommendation of Martijn Pieters.
    soup = BeautifulSoup(resp, from_encoding=resp.info().get_param('charset'), features="html.parser")  
    external_links = set()
    internal_links = set()
    for line in soup.find_all('a'):
        link = line.get('href')
        if not link:
            continue
        if link.startswith('http'):
            external_links.add(link)
        else:
            internal_links.add(link)

    # Depending on usage, full internal links may be preferred.
    full_internal_links = {
        urllib.parse.urljoin(url, internal_link) 
        for internal_link in internal_links
    }

    data = []

    # all unique external and full internal links.
    for link in external_links.union(full_internal_links):
        data.append(link)


    data = deleteduplicates(data)

    if(type=='pdf'):
        data = [x for x in data if x.endswith("pdf")]
    elif(type=='doc'):
        data = [x for x in data if x.endswith("docx")]
        
    
    return data








#Tries to get data from doc files but society has damned windows unfortunately. So non of the libraries work.
def testmethod(url):
    try:
        r = requests.get(url, stream=True)
    except:
        return []

    p = BytesIO(r.content)
    p.seek(0,os.SEEK_END)

    
    tmp_file= tempfile.NamedTemporaryFile(dir=r"C:/Users/88017/Documents/Fall21/SoftEng_Project/cse327.1.5/Snippet-project/temp", suffix=".doc")
    tmp_file.write(p.getvalue())
    tmp_file.seek(0)
    """
    text = textract.process(tmp_file.name)
    text.decode('utf-8')
    text = CONTROL_CHAR_RE.sub('', text)
    tmp_file.close()
    """
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    word.visible = False
    wb = word.Documents.Open(tmp_file.name)
    file= word.ActiveDocument(wb)
    
    data = []
    data.append(file.Range().Text)
    tmp_file.close()

    return data


def givepdf(i):
    try:
        r = requests.get(i).content
    except:
        return "no data"

    
    p = BytesIO(r)
    p.seek(0, os.SEEK_END)

    try:
        read_pdf = PyPDF2.PdfFileReader(p, strict=False)
    except:
        return "no data"
    
    count = read_pdf.numPages
    pages_txt = []

    for i in range(count):
        page = read_pdf.getPage(i)
        pages_txt.append(page.extractText())

    data = ' '.join(pages_txt)

    return data


def givedocx(url):
    try:
        r = requests.get(url).content
    except:
        "bad request"

    p = BytesIO(r)
    p.seek(0, os.SEEK_END)

    try:
        document = docx.Document(p)
    except:
        return "no data"

    
    
    data = []
    for p in document.paragraphs:
        data.append(p.text) 



    data = ' '.join(data)

    return data